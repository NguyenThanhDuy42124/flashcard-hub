"""FastAPI main application with all routes."""
from fastapi import FastAPI, Depends, UploadFile, File, HTTPException, Query, Request, Form, Body
from fastapi.responses import StreamingResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from sqlalchemy import func
from datetime import datetime, timedelta
from typing import List, Optional
from pathlib import Path
import asyncio
import os
import sys
import logging
import json
import html
from io import BytesIO
import math
import random

from docx import Document

# Try to setup Rich logging, fallback to standard logging if not available
try:
    from rich.logging import RichHandler
    from rich.console import Console
    console = Console()
    logging_handler = RichHandler(console=console, rich_tracebacks=True)
    has_rich = True
except ImportError:
    console = None
    logging_handler = logging.StreamHandler()
    has_rich = False

from database import get_db, engine, Base, SessionLocal
from models import User, Deck, Card, CardReview, StudySession
from schemas import (
    CardCreate, CardResponse, DeckCreate, DeckUpdate, DeckResponse,
    UserCreate, UserResponse, CardReviewRequest, StudySessionCreate,
    StudySessionResponse, UserProgressResponse, CreateDeckFromHTMLRequest,
    AppendHTMLRequest, ExamCreateRequest, ExamCreateResponse
)
from parser import parse_html_file
from srs_engine import sm2_engine

# Setup logging with or without Rich colors
logging.basicConfig(
    level=logging.INFO,
    format="%(message)s" if has_rich else "%(asctime)s - %(levelname)s - %(message)s",
    handlers=[logging_handler]
)
logger = logging.getLogger("flashcard_hub")
access_logger = logging.getLogger("uvicorn.access")
access_stream_handler = logging.StreamHandler(stream=sys.stdout)
access_stream_handler.setFormatter(logging.Formatter("%(asctime)s %(message)s"))
access_logger.addHandler(access_stream_handler)
access_logger.setLevel(logging.INFO)
access_logger.propagate = False

TIME_LIMIT_OPTIONS = {
    "1h": timedelta(hours=1),
    "2h": timedelta(hours=2),
    "3h": timedelta(hours=3),
    "1w": timedelta(weeks=1),
    "1m": timedelta(days=30),
    "unlimited": None,
}

EXAM_EXPIRATION_SWEEP_SECONDS = 15 * 60

# Initialize database schema
try:
    from alembic.config import Config
    from alembic.command import upgrade as alembic_upgrade
    from sqlalchemy import text, inspect
    
    def ensure_card_columns_exist():
        """Ensure cards and decks tables have required columns."""
        try:
            inspector = inspect(engine)
            card_columns = [col['name'] for col in inspector.get_columns('cards')]
            deck_columns = [col['name'] for col in inspector.get_columns('decks')]
            
            if 'title' not in card_columns:
                logger.warning("⚠️ Adding missing 'title' column to cards table")
                with engine.connect() as conn:
                    conn.execute(text("ALTER TABLE cards ADD COLUMN title VARCHAR(255)"))
                    conn.commit()
            
            if 'chapter' not in card_columns:
                logger.warning("⚠️ Adding missing 'chapter' column to cards table")
                with engine.connect() as conn:
                    conn.execute(text("ALTER TABLE cards ADD COLUMN chapter VARCHAR(100)"))
                    conn.commit()

            if 'position' not in card_columns:
                logger.warning("⚠️ Adding missing 'position' column to cards table")
                with engine.connect() as conn:
                    conn.execute(text("ALTER TABLE cards ADD COLUMN position INTEGER"))
                    conn.commit()

            if 'allow_card_additions' not in deck_columns:
                logger.warning("⚠️ Adding missing 'allow_card_additions' column to decks table")
                with engine.connect() as conn:
                    conn.execute(text("ALTER TABLE decks ADD COLUMN allow_card_additions BOOLEAN DEFAULT 1"))
                    conn.commit()

            if 'expired_at' not in deck_columns:
                logger.warning("⚠️ Adding missing 'expired_at' column to decks table")
                with engine.connect() as conn:
                    conn.execute(text("ALTER TABLE decks ADD COLUMN expired_at DATETIME"))
                    conn.commit()

            if 'status' not in deck_columns:
                logger.warning("⚠️ Adding missing 'status' column to decks table")
                with engine.connect() as conn:
                    conn.execute(text("ALTER TABLE decks ADD COLUMN status VARCHAR(20) DEFAULT 'active'"))
                    conn.execute(text("UPDATE decks SET status = 'active' WHERE status IS NULL"))
                    conn.commit()
            
            logger.info("✅ Card columns verified")
        except Exception as e:
            logger.error(f"❌ Error checking card columns: {e}")
    
    def run_migrations():
        """Run Alembic migrations automatically on startup."""
        try:
            db_url = os.getenv("DATABASE_URL", f"sqlite:///{Path(__file__).parent.parent / 'flashcard_hub.db'}")
            alembic_config = Config(str(Path(__file__).parent / "alembic.ini"))
            alembic_config.set_main_option("sqlalchemy.url", db_url)
            alembic_upgrade(alembic_config, "head")
            logger.info("✅ Database migrations completed")
            ensure_card_columns_exist()
        except Exception as e:
            logger.error(f"❌ Migration error: {e}")
            logger.info("🔧 Attempting column verification as fallback")
            ensure_card_columns_exist()
    
    def seed_sample_data():
        """Create sample data if database is empty."""
        try:
            db = SessionLocal()
            deck_count = db.query(Deck).count()
            
            if deck_count == 0:
                logger.info("📚 Database is empty, creating sample data...")
                
                # Create test user
                user = User(
                    username="demo_user",
                    email="demo@flashcard.local",
                    password_hash="demo_hash"
                )
                db.add(user)
                db.flush()
                
                # Create sample deck
                deck = Deck(
                    title="Tổng Ôn Mạng Máy Tính",
                    description="Ôn tập kiến thức cơ bản về mạng máy tính",
                    owner_id=user.id,
                    is_public=True,
                    tag="networking"
                )
                db.add(deck)
                db.flush()
                
                # Create sample cards
                sample_cards = [
                    ("TCP/IP là gì?", "TCP/IP là một bộ giao thức mạng"), 
                    ("OSI Model có mấy layer?", "OSI Model có 7 layer"),
                    ("Port 80 dùng cho gì?", "Port 80 dùng cho HTTP"),
                    ("DNS là gì?", "DNS dùng để chuyển đổi tên miền thành IP"),
                    ("DHCP là gì?", "DHCP dùng để cấp phát địa chỉ IP tự động"),
                ]
                
                for front, back in sample_cards:
                    card = Card(
                        deck_id=deck.id,
                        front=front,
                        back=back,
                        title=front[:50],
                        chapter="Chapter 1"
                    )
                    db.add(card)
                
                db.commit()
                logger.info(f"✅ Sample data created: 1 user, 1 deck, {len(sample_cards)} cards")
            else:
                logger.info(f"✅ Database has {deck_count} decks, skipping sample data")
            
            db.close()
        except Exception as e:
            logger.error(f"❌ Error seeding sample data: {e}")
    
    # Run migrations on startup
    run_migrations()
    seed_sample_data()
except ImportError:
    logger.warning("⚠️ Alembic not available, skipping migrations")
    # Fallback: create all tables (only for development)
    Base.metadata.create_all(bind=engine)


def normalize_card_positions(db: Session, deck_id: int):
    """Ensure every card in the deck has a sequential position."""
    cards = db.query(Card).filter(Card.deck_id == deck_id).order_by(Card.position.nulls_last(), Card.created_at).all()
    changed = False
    for idx, card in enumerate(cards, start=1):
        if card.position is None or card.position != idx:
            card.position = idx
            changed = True
    if changed:
        db.commit()


def move_card_to_position(db: Session, card: Card, new_position: int):
    """Reorder card inside its deck, shifting neighbors accordingly."""
    normalize_card_positions(db, card.deck_id)

    total_cards = db.query(Card).filter(Card.deck_id == card.deck_id).count()
    target_pos = max(1, min(new_position, total_cards))

    if card.position is None:
        card.position = total_cards

    if target_pos == card.position:
        return

    if target_pos < card.position:
        db.query(Card).filter(
            Card.deck_id == card.deck_id,
            Card.position >= target_pos,
            Card.position < card.position,
            Card.id != card.id
        ).update({Card.position: Card.position + 1}, synchronize_session=False)
    else:
        db.query(Card).filter(
            Card.deck_id == card.deck_id,
            Card.position <= target_pos,
            Card.position > card.position,
            Card.id != card.id
        ).update({Card.position: Card.position - 1}, synchronize_session=False)

    card.position = target_pos
    db.commit()


def parse_quiz_meta(back_content: str):
    """Extract quiz metadata from card back if it is quiz formatted."""
    if not back_content or not isinstance(back_content, str):
        return None
    if not back_content.startswith("__QUIZ__::"):
        return None
    try:
        payload = json.loads(back_content.replace("__QUIZ__::", "", 1))
        if payload.get("type") == "quiz":
            return payload
    except Exception:
        return None
    return None


def export_deck_to_html(deck: Deck, cards: List[Card]) -> str:
    """Render deck and cards into a standalone HTML string for download."""
    card_blocks = []
    for idx, card in enumerate(cards, start=1):
        quiz_meta = parse_quiz_meta(card.back)
        safe_title = html.escape(card.title or card.front or f"Thẻ {idx}")
        safe_front = html.escape(card.front or "")
        safe_chapter = html.escape(card.chapter or "Chưa phân chương")
        body_html = ""

        if quiz_meta:
            options = quiz_meta.get("options", {})
            correct_badge = '<span style="color:#16a34a;font-weight:700">(Đúng)</span>'
            options_html = "".join([
                f"<li><strong>{html.escape(str(k))}.</strong> {html.escape(str(v))}"
                f" {correct_badge if k == quiz_meta.get('correct') else ''}</li>"
                for k, v in options.items()
            ])
            explanation = html.escape(quiz_meta.get("explanation") or "")
            body_html = (
                f"<p><strong>Dạng:</strong> Trắc nghiệm</p>"
                f"<ul>{options_html}</ul>"
                f"<p><strong>Giải thích:</strong> {explanation or 'Chưa có'}</p>"
            )
        else:
            back_lines = [line.strip() for line in (card.back or "").splitlines() if line.strip()]
            if len(back_lines) > 1:
                body_html = "<ul>" + "".join(
                    [f"<li>{html.escape(line)}</li>" for line in back_lines]
                ) + "</ul>"
            else:
                body_html = f"<p>{html.escape(card.back or '')}</p>"

        card_blocks.append(
            f"""
            <section style=\"padding:16px;border:1px solid #e5e7eb;border-radius:12px;background:#fff;box-shadow:0 8px 24px -12px rgba(0,0,0,0.15);\">
                <div style=\"display:flex;justify-content:space-between;align-items:center;margin-bottom:8px;color:#475569;font-size:14px;\">
                    <span style=\"font-weight:700;\">#{idx}</span>
                    <span style=\"padding:4px 10px;background:#eef2ff;color:#4338ca;border-radius:999px;font-weight:600;\">{safe_chapter}</span>
                </div>
                <h3 style=\"margin:6px 0;color:#0f172a;font-size:20px;font-weight:800;\">{safe_title}</h3>
                <p style=\"margin:10px 0 12px;color:#1e293b;font-size:16px;line-height:1.6;\"><strong>Q:</strong> {safe_front}</p>
                <div style=\"color:#0f172a;font-size:15px;line-height:1.6;\">{body_html}</div>
            </section>
            """
        )

    cards_html = "\n".join(card_blocks)
    safe_deck_title = html.escape(deck.title or "Deck")
    safe_description = html.escape(deck.description or "")

    return f"""<!DOCTYPE html>
<html lang=\"vi\">
<head>
    <meta charset=\"UTF-8\" />
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\" />
    <title>{safe_deck_title} - Export</title>
    <style>
        body {{ font-family: 'Inter', sans-serif; background:#f8fafc; padding:24px; color:#0f172a; }}
        .deck-header {{ margin-bottom:24px; }}
        .deck-header h1 {{ font-size:28px; margin:0; }}
        .deck-header p {{ color:#475569; margin-top:8px; }}
        .grid {{ display:grid; gap:16px; grid-template-columns:repeat(auto-fit, minmax(260px, 1fr)); }}
    </style>
</head>
<body>
    <div class=\"deck-header\">
        <h1>📚 {safe_deck_title}</h1>
        <p>{safe_description}</p>
        <p style=\"color:#475569;\">Tổng số thẻ: {len(cards)}</p>
    </div>
    <div class=\"grid\">{cards_html}</div>
</body>
</html>"""


def export_deck_to_docx(deck: Deck, cards: List[Card]) -> BytesIO:
    """Render deck and cards into a DOCX file buffer."""
    document = Document()
    document.add_heading(deck.title or "Deck", 0)
    if deck.description:
        document.add_paragraph(deck.description)
    document.add_paragraph(f"Tổng số thẻ: {len(cards)}")

    for idx, card in enumerate(cards, start=1):
        document.add_heading(f"Thẻ #{idx}: {card.title or card.front or 'Không tiêu đề'}", level=1)
        document.add_paragraph(f"Chương: {card.chapter or 'Chưa phân chương'}")
        document.add_paragraph("Câu Hỏi:")
        document.add_paragraph(card.front or "")

        quiz_meta = parse_quiz_meta(card.back)
        if quiz_meta:
            document.add_paragraph("Câu Trả Lời:")
            document.add_paragraph("Dạng: Trắc nghiệm")
            options = quiz_meta.get("options", {})
            correct_key = quiz_meta.get("correct")
            for key, text in options.items():
                label = "(Đúng)" if key == correct_key else "(Sai)"
                document.add_paragraph(f"{key}. {text} {label}")
            if quiz_meta.get("explanation"):
                document.add_paragraph(f"Giải thích: {quiz_meta.get('explanation')}")
        else:
            document.add_paragraph("Câu Trả Lời:")
            if card.back:
                for line in card.back.splitlines():
                    document.add_paragraph(line)
            else:
                document.add_paragraph("")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# Initialize FastAPI app
app = FastAPI(
    title="Flashcard Hub API",
    description="Full-stack flashcard learning platform with SRS",
    version="1.0.0"
)


@app.on_event("startup")
async def start_exam_expiration_job():
    """Start periodic background task for exam expiration cleanup."""
    asyncio.create_task(exam_expiration_sweeper())
    logger.info("⏱️ Exam expiration sweeper started (every 15 minutes)")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3001", "http://localhost:8000", "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Add response headers for insecure context
@app.middleware("http")
async def add_secure_headers(request, call_next):
    response = await call_next(request)
    response.headers["Cross-Origin-Resource-Policy"] = "cross-origin"
    response.headers["Cross-Origin-Embedder-Policy"] = "require-corp"
    return response


# Minimal request logger: keep only method/path and client IP
@app.middleware("http")
async def log_request_meta(request: Request, call_next):
    forwarded_for = request.headers.get("x-forwarded-for")
    client_ip = forwarded_for.split(",")[0].strip() if forwarded_for else (request.client.host if request.client else "unknown")

    try:
        response = await call_next(request)
        return response
    finally:
        msg = f"🌐 {request.method} {request.url.path} | ip={client_ip}"
        # Send to uvicorn access logger so it shows up in terminal
        access_logger.info(msg)
        print(msg, flush=True)


def select_cards_for_deck(cards: List[Card], count: int, scope: str = "deck") -> List[Card]:
    """Randomly select cards with optional chapter balancing."""
    if count <= 0 or not cards:
        return []

    if scope == "chapter":
        grouped = {}
        for card in cards:
            grouped.setdefault(card.chapter or "Khác", []).append(card)

        selected: List[Card] = []
        # Round-robin pull across chapters to stay balanced
        while grouped and len(selected) < min(count, sum(len(v) for v in grouped.values())):
            for chapter in list(grouped.keys()):
                bucket = grouped.get(chapter, [])
                if not bucket:
                    grouped.pop(chapter, None)
                    continue
                pick = bucket.pop(random.randrange(len(bucket)))
                selected.append(pick)
                if not bucket:
                    grouped.pop(chapter, None)
                if len(selected) >= count:
                    break

        # If we still need more (deck có ít thẻ), allow repeats to fill
        if len(selected) < count:
            selected.extend(random.choices(cards, k=count - len(selected)))

        return selected[:count]

    pool = list(cards)
    random.shuffle(pool)
    if len(pool) >= count:
        return pool[:count]

    # Not enough unique cards -> allow repeats to reach the target count
    result = pool[:]
    result.extend(random.choices(pool, k=count - len(result)))
    return result[:count]


def is_deck_expired(deck: Deck, now: Optional[datetime] = None) -> bool:
    """Check whether a deck has passed expiration time."""
    if not deck.expired_at:
        return False
    now_time = now or datetime.utcnow()
    return deck.expired_at <= now_time


def mark_deck_expired_if_needed(db: Session, deck: Deck, now: Optional[datetime] = None) -> bool:
    """Soft-expire a deck if needed. Returns True if updated."""
    if deck.status == "expired":
        return False
    if is_deck_expired(deck, now=now):
        deck.status = "expired"
        return True
    return False


def expire_due_exams(db: Session, now: Optional[datetime] = None) -> int:
    """Soft-expire all due exam decks and return affected count."""
    now_time = now or datetime.utcnow()
    due_decks = db.query(Deck).filter(
        Deck.expired_at.isnot(None),
        Deck.expired_at <= now_time,
        (Deck.status.is_(None)) | (Deck.status != "expired")
    ).all()

    for deck in due_decks:
        deck.status = "expired"

    if due_decks:
        db.commit()

    return len(due_decks)


async def exam_expiration_sweeper():
    """Periodic job to expire exams every 15 minutes."""
    while True:
        await asyncio.sleep(EXAM_EXPIRATION_SWEEP_SECONDS)
        db = SessionLocal()
        try:
            changed = expire_due_exams(db)
            if changed > 0:
                logger.info(f"🧹 Expiration sweep marked {changed} deck(s) as expired")

            if os.getenv("DELETE_EXPIRED_EXAMS", "false").lower() == "true":
                now_time = datetime.utcnow()
                hard_deleted = db.query(Deck).filter(
                    Deck.expired_at.isnot(None),
                    Deck.expired_at <= now_time,
                    Deck.status == "expired"
                ).delete(synchronize_session=False)
                if hard_deleted > 0:
                    db.commit()
                    logger.info(f"🗑️ Hard deleted {hard_deleted} expired deck(s)")
        except Exception as e:
            db.rollback()
            logger.error(f"❌ Expiration sweep failed: {e}")
        finally:
            db.close()


# ============== DECK ENDPOINTS ==============

@app.get("/api/decks", response_model=List[DeckResponse])
async def list_decks(
    db: Session = Depends(get_db),
    skip: int = Query(0, ge=0),
    limit: int = Query(10, ge=1, le=100),
    tag: str = Query(None),
    status: Optional[str] = Query(None)
):
    """List all public decks with pagination and filtering."""
    try:
        logger.info(f"📋 Fetching decks: skip={skip}, limit={limit}, tag={tag}")
        
        expire_due_exams(db)

        query = db.query(Deck).filter(Deck.is_public == True)

        if tag:
            query = query.filter(Deck.tag == tag)

        if status == "expired":
            query = query.filter(Deck.status == "expired")
        elif status == "active":
            query = query.filter((Deck.status.is_(None)) | (Deck.status == "active"))
        elif status is None:
            query = query.filter((Deck.status.is_(None)) | (Deck.status != "expired"))

        decks = query.offset(skip).limit(limit).all()
        logger.info(f"✅ Found {len(decks)} public decks")
        
        return decks
    except Exception as e:
        logger.error(f"❌ Error loading decks: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error loading decks: {str(e)}")


@app.post("/api/decks/create", response_model=DeckResponse)
async def create_deck(
    deck_data: DeckCreate,
    db: Session = Depends(get_db),
    user_id: int = 1  # In production, get from JWT token
):
    """Create a new deck manually via UI."""
    # Create deck
    new_deck = Deck(
        title=deck_data.title,
        description=deck_data.description,
        owner_id=user_id,
        tag=deck_data.tag,
        is_public=deck_data.is_public,
        allow_card_additions=deck_data.allow_card_additions
    )
    db.add(new_deck)
    db.flush()

    # Add cards if provided
    for card_data in deck_data.cards:
        new_card = Card(
            deck_id=new_deck.id,
            front=card_data.front,
            back=card_data.back
        )
        db.add(new_card)

    db.commit()
    db.refresh(new_deck)
    return new_deck


@app.post("/api/decks/upload-html", response_model=DeckResponse)
async def upload_html_deck(
    file: UploadFile = File(...),
    title: str = Form(None),
    description: str = Form(None),
    db: Session = Depends(get_db),
    user_id: int = 1
):
    """Upload and parse an HTML file into a database deck."""
    try:
        print(f"📥 Uploading file: {file.filename}")
        
        # Read file content
        content = await file.read()
        print(f"📄 File size: {len(content)} bytes")
        
        html_content = content.decode('utf-8')
        print(f"✅ HTML decoded, length: {len(html_content)}")

        # Parse HTML
        print("🔍 Parsing HTML...")
        parsed_data = parse_html_file(html_content)
        print(f"✅ Parsed successfully: title='{parsed_data['title']}', cards={len(parsed_data['cards'])}")

        # Create deck
        new_deck = Deck(
            title=title or parsed_data['title'],
            description=description if description is not None else parsed_data['description'],
            owner_id=user_id,
            tag=parsed_data.get('tag'),
            is_public=True
        )
        db.add(new_deck)
        db.flush()
        print(f"✅ Deck created: id={new_deck.id}")

        # Add cards from parsed data
        for idx, card_data in enumerate(parsed_data['cards']):
            new_card = Card(
                deck_id=new_deck.id,
                front=card_data['front'],
                back=card_data['back'],
                title=card_data.get('title'),
                chapter=card_data.get('chapter')
            )
            db.add(new_card)

        db.commit()
        db.refresh(new_deck)
        print(f"✅ Deck committed with {len(parsed_data['cards'])} cards")

        return new_deck

    except Exception as e:
        print(f"❌ Error in upload_html_deck: {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Failed to parse HTML: {str(e)}")


@app.post("/api/decks/create-from-html-content", response_model=DeckResponse)
async def create_deck_from_html_content(
    request: CreateDeckFromHTMLRequest,
    db: Session = Depends(get_db),
    user_id: int = 1
):
    """Create a deck from HTML content (paste from text)."""
    try:
        # Parse HTML content
        parsed_data = parse_html_file(request.html_content)

        # Create deck with provided title and description
        new_deck = Deck(
            title=request.title if request.title else parsed_data.get('title', 'Untitled Deck'),
            description=request.description if request.description is not None else parsed_data.get('description', ''),
            owner_id=user_id,
            tag=request.tag if request.tag is not None else parsed_data.get('tag'),
            is_public=True
        )
        db.add(new_deck)
        db.flush()

        # Add cards from parsed data
        for card_data in parsed_data['cards']:
            new_card = Card(
                deck_id=new_deck.id,
                front=card_data['front'],
                back=card_data['back'],
                title=card_data.get('title'),
                chapter=card_data.get('chapter')
            )
            db.add(new_card)

        db.commit()
        db.refresh(new_deck)

        return new_deck

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Failed to parse HTML: {str(e)}")


@app.get("/api/decks/{deck_id}", response_model=DeckResponse)
async def get_deck(deck_id: int, db: Session = Depends(get_db)):
    """Get a specific deck with all its cards."""
    try:
        logger.info(f"📖 Fetching deck: deck_id={deck_id}")
        deck = db.query(Deck).filter(Deck.id == deck_id).first()

        if not deck:
            logger.warning(f"⚠️ Deck not found: deck_id={deck_id}")
            all_decks = db.query(Deck).count()
            logger.info(f"   Available decks in database: {all_decks}")
            raise HTTPException(status_code=404, detail=f"Deck {deck_id} not found")

        if mark_deck_expired_if_needed(db, deck):
            db.commit()
            db.refresh(deck)

        if deck.status == "expired":
            raise HTTPException(status_code=410, detail="Đề đã hết hạn")

        logger.info(f"✅ Deck found: {deck.title} (ID: {deck_id}, Cards: {len(deck.cards)})")
        return deck
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Error fetching deck: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/decks/{deck_id}/export")
async def export_deck(
    deck_id: int,
    format: str = Query("html", regex="^(html|docx)$", description="Định dạng xuất: html hoặc docx"),
    sort_by: str = Query("chapter", description="Sort by: position, chapter, title, created"),
    chapters: str = Query(None, description="Filter chapters, phân tách bằng dấu phẩy"),
    db: Session = Depends(get_db)
):
    """Export a deck to HTML or DOCX for offline use or re-upload."""
    deck = db.query(Deck).filter(Deck.id == deck_id).first()
    if not deck:
        raise HTTPException(status_code=404, detail=f"Deck {deck_id} not found")

    query = db.query(Card).filter(Card.deck_id == deck_id)

    if chapters:
        chapter_list = [c.strip() for c in chapters.split(",") if c.strip() and c.strip() != "Tất cả"]
        if chapter_list:
            query = query.filter(Card.chapter.in_(chapter_list))

    if sort_by == "title":
        cards = query.order_by(Card.title, Card.position.nulls_last()).all()
    elif sort_by == "created":
        cards = query.order_by(Card.created_at.desc()).all()
    elif sort_by == "chapter":
        cards = query.order_by(Card.chapter, Card.position.nulls_last(), Card.created_at).all()
    else:  # position is default
        cards = query.order_by(Card.position.nulls_last(), Card.created_at).all()

    filename_base = f"deck-{deck_id}"
    if format == "html":
        html_content = export_deck_to_html(deck, cards)
        return Response(
            content=html_content,
            media_type="text/html; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename={filename_base}.html"}
        )

    buffer = export_deck_to_docx(deck, cards)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename_base}.docx"}
    )


@app.get("/api/decks/{deck_id}/cards", response_model=List[CardResponse])
async def get_deck_cards(
    deck_id: int,
    sort_by: str = Query("chapter", description="Sort by: position, chapter, title, created"),
    chapter: str = Query(None, description="Filter by chapter"),
    search: str = Query(None, description="Search by title"),
    db: Session = Depends(get_db)
):
    """Get all cards for a specific deck with optional sorting and filtering."""
    try:
        logger.info(f"🃏 Fetching cards: deck_id={deck_id}, sort_by={sort_by}, chapter={chapter}, search={search}")
        
        deck = db.query(Deck).filter(Deck.id == deck_id).first()

        if not deck:
            logger.warning(f"⚠️ Deck not found for cards: deck_id={deck_id}")
            all_decks = db.query(Deck).count()
            logger.info(f"   Available decks in database: {all_decks}")
            raise HTTPException(status_code=404, detail=f"Deck {deck_id} not found")

        if mark_deck_expired_if_needed(db, deck):
            db.commit()
            db.refresh(deck)

        if deck.status == "expired":
            raise HTTPException(status_code=410, detail="Đề đã hết hạn")

        query = db.query(Card).filter(Card.deck_id == deck_id)
        
        # Filter by chapter if provided
        if chapter and chapter != "Tất cả":
            query = query.filter(Card.chapter == chapter)
        
        # Search by title if provided
        if search:
            query = query.filter(Card.title.ilike(f"%{search}%"))
        
        # Sort results
        if sort_by == "title":
            cards = query.order_by(Card.title, Card.position.nulls_last()).all()
        elif sort_by == "created":
            cards = query.order_by(Card.created_at.desc()).all()
        elif sort_by == "chapter":
            cards = query.order_by(Card.chapter, Card.position.nulls_last(), Card.created_at).all()
        else:  # position is default
            cards = query.order_by(Card.position.nulls_last(), Card.created_at).all()
        
        logger.info(f"✅ Found {len(cards)} cards for deck '{deck.title}'")
        return cards
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Error fetching cards: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/decks/{deck_id}", response_model=DeckResponse)
async def update_deck(
    deck_id: int,
    deck_data: DeckUpdate,
    db: Session = Depends(get_db),
    user_id: int = 1
):
    """Update a deck."""
    deck = db.query(Deck).filter(Deck.id == deck_id).first()

    if not deck:
        raise HTTPException(status_code=404, detail="Deck not found")

    if deck.owner_id != user_id:
        raise HTTPException(status_code=403, detail="Not authorized")

    # Update fields
    if deck_data.title is not None:
        deck.title = deck_data.title
    if deck_data.description is not None:
        deck.description = deck_data.description
    if deck_data.tag is not None:
        deck.tag = deck_data.tag
    if deck_data.is_public is not None:
        deck.is_public = deck_data.is_public
    if deck_data.allow_card_additions is not None:
        deck.allow_card_additions = deck_data.allow_card_additions

    deck.updated_at = datetime.utcnow()

    db.commit()
    db.refresh(deck)
    return deck


@app.delete("/api/decks/{deck_id}")
async def delete_deck(
    deck_id: int,
    db: Session = Depends(get_db),
    user_id: int = 1
):
    """Delete a deck."""
    deck = db.query(Deck).filter(Deck.id == deck_id).first()

    if not deck:
        raise HTTPException(status_code=404, detail="Deck not found")

    if deck.owner_id != user_id:
        raise HTTPException(status_code=403, detail="Not authorized")

    db.delete(deck)
    db.commit()

    return {"message": "Deck deleted successfully"}



@app.post("/api/decks/{deck_id}/append-from-html")
async def append_cards_from_html(
    deck_id: int,
    request: AppendHTMLRequest,
    db: Session = Depends(get_db),
    admin_override: bool = Query(False)
):
    try:
        from parser import parse_html_file
        parsed_data = parse_html_file(request.html_content)
        
        deck = db.query(Deck).filter(Deck.id == deck_id).first()
        if not deck:
            raise HTTPException(status_code=404, detail="Deck not found")

        if not deck.allow_card_additions and not admin_override:
            raise HTTPException(status_code=403, detail="Deck đã khoá thêm thẻ")
        # Nếu deck chưa có tag, gán theo loại phát hiện từ nội dung
        if not deck.tag and parsed_data.get('tag'):
            deck.tag = parsed_data.get('tag')
            
        added_cards = []
        # Determine current max position in this deck
        max_pos = db.query(func.max(Card.position)).filter(Card.deck_id == deck_id).scalar()
        next_pos = (max_pos or 0) + 1

        for card_data in parsed_data['cards']:
            new_card = Card(
                deck_id=deck_id,
                front=card_data['front'],
                back=card_data['back'],
                title=card_data.get('title'),
                chapter=card_data.get('chapter'),
                position=card_data.get('position') or next_pos
            )
            next_pos += 1
            db.add(new_card)
            added_cards.append(new_card)
            
        db.commit()
        return {"message": f"Added {len(added_cards)} cards successfully"}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/exams/create", response_model=ExamCreateResponse)
async def create_exam_deck(
    exam_data: ExamCreateRequest,
    db: Session = Depends(get_db),
    user_id: int = 1
):
    """Generate an exam deck by mixing questions from multiple decks using percentages."""
    if not exam_data.selections:
        raise HTTPException(status_code=400, detail="Chưa chọn deck nào")

    if exam_data.total_questions <= 0:
        raise HTTPException(status_code=400, detail="Tổng số câu phải > 0")

    if exam_data.random_scope not in ("deck", "chapter"):
        raise HTTPException(status_code=400, detail="random_scope không hợp lệ")

    valid_selections = [s for s in exam_data.selections if s.percentage > 0]
    if not valid_selections:
        raise HTTPException(status_code=400, detail="Phần trăm phải lớn hơn 0")

    percent_sum = sum(sel.percentage for sel in valid_selections)
    if percent_sum <= 0:
        raise HTTPException(status_code=400, detail="Tổng phần trăm không hợp lệ")

    # Normalize distribution and allocate leftover by fractional part
    distribution = []
    for sel in valid_selections:
        weight = sel.percentage / percent_sum
        ideal = weight * exam_data.total_questions
        base = math.floor(ideal)
        frac = ideal - base
        distribution.append({"deck_id": sel.deck_id, "count": base, "frac": frac})

    remainder = exam_data.total_questions - sum(item["count"] for item in distribution)
    distribution.sort(key=lambda x: x["frac"], reverse=True)
    for i in range(remainder):
        distribution[i % len(distribution)]["count"] += 1

    selected_cards: List[Card] = []
    all_cards_pool: List[Card] = []

    for item in distribution:
        deck = db.query(Deck).filter(Deck.id == item["deck_id"]).first()
        if not deck:
            raise HTTPException(status_code=404, detail=f"Deck {item['deck_id']} không tồn tại")

        cards = db.query(Card).filter(Card.deck_id == deck.id).all()
        all_cards_pool.extend(cards)

        if item["count"] <= 0:
            continue

        selected_cards.extend(
            select_cards_for_deck(cards, item["count"], exam_data.random_scope)
        )

    # Fallback: if not enough unique cards, top up from pooled cards
    if len(selected_cards) < exam_data.total_questions and all_cards_pool:
        deficit = exam_data.total_questions - len(selected_cards)
        selected_cards.extend(random.choices(all_cards_pool, k=deficit))

    # Shuffle to mix decks together
    random.shuffle(selected_cards)

    time_delta = None
    if exam_data.time_limit:
        if exam_data.time_limit not in TIME_LIMIT_OPTIONS:
            raise HTTPException(status_code=400, detail="Thời gian không hợp lệ")
        time_delta = TIME_LIMIT_OPTIONS.get(exam_data.time_limit)

    expires_at: Optional[datetime] = None
    if time_delta:
        expires_at = datetime.utcnow() + time_delta

    description = exam_data.description or "Đề thi trộn từ nhiều deck"
    if exam_data.time_limit and exam_data.time_limit != "unlimited":
        description = f"{description} | Thời gian: {exam_data.time_limit}"

    target_tag = (exam_data.tag or "Quiz").strip() if isinstance(exam_data.tag, str) else "Quiz"
    if target_tag not in ["Quiz", "Flashcard"]:
        raise HTTPException(status_code=400, detail="tag phải là 'Quiz' hoặc 'Flashcard'")

    exam_deck = Deck(
        title=exam_data.title,
        description=description,
        owner_id=user_id,
        tag=target_tag,
        is_public=True,
        allow_card_additions=False,
        expired_at=expires_at,
        status="active"
    )
    db.add(exam_deck)
    db.flush()

    for card in selected_cards:
        new_card = Card(
            deck_id=exam_deck.id,
            title=card.title,
            chapter=card.chapter,
            front=card.front,
            back=card.back,
            position=None
        )
        db.add(new_card)

    db.commit()
    db.refresh(exam_deck)

    return ExamCreateResponse(
        exam_deck_id=exam_deck.id,
        deck_title=exam_deck.title,
        total_questions=len(selected_cards),
        expires_at=expires_at,
        status=exam_deck.status,
        tag=target_tag,
    )


# ============== CARD ENDPOINTS ==============

@app.post("/api/cards", response_model=CardResponse)
async def create_card(
    deck_id: int,
    card_data: CardCreate,
    db: Session = Depends(get_db),
    admin_override: bool = Query(False)
):
    """Create a new card in a deck with title and chapter."""
    deck = db.query(Deck).filter(Deck.id == deck_id).first()

    if not deck:
        raise HTTPException(status_code=404, detail="Deck not found")

    if not deck.allow_card_additions and not admin_override:
        raise HTTPException(status_code=403, detail="Deck đã khoá thêm thẻ")

    max_pos = db.query(func.max(Card.position)).filter(Card.deck_id == deck_id).scalar()
    next_pos = (max_pos or 0) + 1

    new_card = Card(
        deck_id=deck_id,
        title=card_data.title,
        chapter=card_data.chapter,
        front=card_data.front,
        back=card_data.back,
        position=card_data.position or next_pos
    )
    db.add(new_card)
    db.commit()
    db.refresh(new_card)

    return new_card


@app.put("/api/cards/{card_id}", response_model=CardResponse)
async def update_card(
    card_id: int,
    card_data: CardCreate,
    db: Session = Depends(get_db)
):
    """Update a card."""
    card = db.query(Card).filter(Card.id == card_id).first()

    if not card:
        raise HTTPException(status_code=404, detail="Card not found")

    card.front = card_data.front
    card.back = card_data.back
    card.title = card_data.title
    card.chapter = card_data.chapter
    if card_data.position:
        move_card_to_position(db, card, int(card_data.position))
    card.updated_at = datetime.utcnow()

    db.commit()
    db.refresh(card)

    return card


@app.post("/api/cards/{card_id}/reorder", response_model=CardResponse)
async def reorder_card(card_id: int, position: int = Body(..., embed=True), db: Session = Depends(get_db)):
    """Move a card to a new position and shift others."""
    card = db.query(Card).filter(Card.id == card_id).first()

    if not card:
        raise HTTPException(status_code=404, detail="Card not found")

    move_card_to_position(db, card, int(position))
    db.refresh(card)
    return card


@app.delete("/api/cards/{card_id}")
async def delete_card(card_id: int, db: Session = Depends(get_db)):
    """Delete a card."""
    card = db.query(Card).filter(Card.id == card_id).first()

    if not card:
        raise HTTPException(status_code=404, detail="Card not found")

    db.delete(card)
    db.commit()

    return {"message": "Card deleted successfully"}


# ============== STUDY & SRS ENDPOINTS ==============

@app.post("/api/study-sessions", response_model=StudySessionResponse)
async def start_study_session(
    session_data: StudySessionCreate,
    db: Session = Depends(get_db),
    user_id: int = 1
):
    """Start a new study session."""
    try:
        logger.info(f"📚 Starting study session for deck_id={session_data.deck_id}, user_id={user_id}")
        
        deck = db.query(Deck).filter(Deck.id == session_data.deck_id).first()

        if not deck:
            logger.warning(f"⚠️ Deck not found: deck_id={session_data.deck_id}")
            raise HTTPException(status_code=404, detail=f"Deck {session_data.deck_id} not found")

        new_session = StudySession(
            user_id=user_id,
            deck_id=session_data.deck_id
        )
        db.add(new_session)
        db.commit()
        db.refresh(new_session)
        
        logger.info(f"✅ Study session created: session_id={new_session.id}")

        return new_session
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Error starting study session: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/cards/{card_id}/review")
async def submit_card_review(
    card_id: int,
    review_data: CardReviewRequest,
    db: Session = Depends(get_db)
):
    """Submit a card review and update SRS data."""
    try:
        card = db.query(Card).filter(Card.id == card_id).first()

        if not card:
            raise HTTPException(status_code=404, detail="Card not found")

        study_session = db.query(StudySession).filter(
            StudySession.id == review_data.study_session_id
        ).first()

        if not study_session:
            raise HTTPException(status_code=404, detail="Study session not found")

        # Get or create review record
        existing_review = db.query(CardReview).filter(
            CardReview.card_id == card_id,
            CardReview.study_session_id == review_data.study_session_id
        ).first()

        if existing_review:
            # Update existing review
            ease_factor = existing_review.ease_factor
            repetitions = existing_review.repetitions
        else:
            # New review
            ease_factor = 2.5
            repetitions = 0

        # Calculate new interval using SM-2
        interval, new_ease, new_repetitions = sm2_engine.calculate_interval(
            repetitions, ease_factor, review_data.quality
        )

        next_review = sm2_engine.get_next_review_date(interval)

        # Create or update review record
        if existing_review:
            existing_review.quality = review_data.quality
            existing_review.ease_factor = new_ease
            existing_review.interval = interval
            existing_review.repetitions = new_repetitions
            existing_review.next_review_date = next_review
            existing_review.reviewed_at = datetime.utcnow()
        else:
            new_review = CardReview(
                card_id=card_id,
                study_session_id=review_data.study_session_id,
                quality=review_data.quality,
                ease_factor=new_ease,
                interval=interval,
                repetitions=new_repetitions,
                next_review_date=next_review
            )
            db.add(new_review)

        # Update study session stats
        study_session.cards_reviewed += 1
        if review_data.quality >= 3:
            study_session.cards_correct += 1

        db.commit()

        return {
            "message": "Review submitted",
            "next_review_date": next_review,
            "ease_factor": new_ease,
            "interval": interval
        }

    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/users/progress", response_model=UserProgressResponse)
async def get_user_progress(
    db: Session = Depends(get_db),
    user_id: int = 1
):
    """Get current user's study statistics."""
    # Total cards reviewed
    total_reviews = db.query(CardReview).join(
        StudySession
    ).filter(
        StudySession.user_id == user_id
    ).count()

    # Total study sessions
    total_sessions = db.query(StudySession).filter(
        StudySession.user_id == user_id
    ).count()

    # Average accuracy
    sessions = db.query(StudySession).filter(
        StudySession.user_id == user_id
    ).all()

    total_correct = sum(s.cards_correct for s in sessions)
    average_accuracy = (total_correct / total_reviews * 100) if total_reviews > 0 else 0

    # Decks studied
    decks_studied = len(set(s.deck_id for s in sessions))

    # Current streak (last session within 24 hours)
    latest_session = db.query(StudySession).filter(
        StudySession.user_id == user_id
    ).order_by(StudySession.started_at.desc()).first()

    current_streak = 1 if latest_session and (
        datetime.utcnow() - latest_session.started_at < timedelta(days=1)
    ) else 0

    return {
        "total_cards_reviewed": total_reviews,
        "total_sessions": total_sessions,
        "average_accuracy": round(average_accuracy, 2),
        "decks_studied": decks_studied,
        "current_streak": current_streak
    }


@app.get("/api/health")
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat()}


@app.get("/api/server/status")
async def server_status(db: Session = Depends(get_db)):
    """Get detailed server status with database information."""
    try:
        logger.info("📊 Server status requested")
        
        # Get database statistics
        total_decks = db.query(Deck).count()
        total_cards = db.query(Card).count()
        total_users = db.query(User).count()
        public_decks = db.query(Deck).filter(Deck.is_public == True).count()
        
        status_data = {
            "status": "operational",
            "timestamp": datetime.utcnow().isoformat(),
            "database": {
                "total_users": total_users,
                "total_decks": total_decks,
                "public_decks": public_decks,
                "total_cards": total_cards,
            },
            "version": "1.0.0"
        }
        
        logger.info(
            f"✅ Server status: {total_decks} decks, "
            f"{total_cards} cards, {total_users} users"
        )
        
        return status_data
    except Exception as e:
        logger.error(f"❌ Error getting server status: {e}")
        raise HTTPException(status_code=500, detail=f"Error getting server status: {str(e)}")


# Mount frontend React build - serve static files (at END to not block /api routes!)
# StaticFiles with html=True will serve index.html for any non-existent routes (SPA fallback)
# Use absolute path since app.py chdir to backend

# Get the project root (parent of backend directory)
backend_dir = Path(__file__).parent
project_root = backend_dir.parent
frontend_build = project_root / "frontend" / "build"

print(f"🔍 Backend dir: {backend_dir}")
print(f"🔍 Project root: {project_root}")
print(f"🔍 Frontend build path: {frontend_build}")
print(f"🔍 Frontend build exists: {frontend_build.exists()}")
if frontend_build.exists():
    print(f"🔍 Frontend build contents: {list(frontend_build.iterdir())[:5]}")
    
    from fastapi.responses import FileResponse
    import mimetypes
    
    # Static files specifically for JS/CSS assets
    static_dir = frontend_build / "static"
    if static_dir.exists():
        app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")
    
    # SPA Catch-all exceptions handler for cleaner 404s
    from fastapi.responses import JSONResponse
    
    @app.exception_handler(404)
    async def custom_404_handler(request: Request, exc: HTTPException):
        if request.method == "GET" and not request.url.path.startswith("/api/"):
            index_path = frontend_build / "index.html"
            if index_path.exists():
                return FileResponse(str(index_path), media_type="text/html")
        return JSONResponse(status_code=404, content={"detail": getattr(exc, "detail", "Not Found")})

    # Catch-all route for SPA
    @app.get("/{full_path:path}", include_in_schema=False)
    async def serve_spa(request: Request, full_path: str):
        if full_path.startswith("api/"):
            raise HTTPException(status_code=404, detail="Not Found")

        # Fallback to files in the root of build
        file_path = frontend_build / full_path
        if file_path.is_file():
            # Get the correct MIME type
            mime_type, _ = mimetypes.guess_type(str(file_path))
            return FileResponse(str(file_path), media_type=mime_type)

        index_path = frontend_build / "index.html"
        if index_path.exists():
            return FileResponse(str(index_path), media_type="text/html")
        raise HTTPException(status_code=404, detail="Index.html not found")

    print(f"✅ Frontend mounted successfully at /")
else:
    print(f"❌ Frontend build not found at {frontend_build}")
